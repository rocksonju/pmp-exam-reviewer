import re
import io
import os
import signal
import csv
from flask import Flask, request, send_file, render_template, jsonify

# DOCX
from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# PDF
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, PageBreak, Paragraph
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.colors import HexColor

app = Flask(__name__)

# ─── Parser ───────────────────────────────────────────────────────────────────

def parse_markdown(content: str):
    blocks = re.split(r"(?=^Question No\.)", content, flags=re.MULTILINE)
    results = []

    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.splitlines()

        q_no = ""
        q_id = ""
        question_text = []
        choices = []
        selected = []
        process_group = []
        task_area = []
        justification = []
        domain = []
        task_domain = []
        reference = []
        state = "question"

        for line in lines:
            line_str = line.strip()
            if not line_str and state not in ["question", "justification"]:
                continue

            m = re.match(r"^Question No\.\s*(.+?)\s*-\s*ID\s*:\s*(.+)", line_str)
            if m:
                q_no = m.group(1).strip()
                q_id = m.group(2).strip()
                state = "question"
                continue

            if re.match(r"^Choice\s*\d+", line_str):
                state = "choice"
                continue

            if line_str == "Selected Choice:":
                state = "selected"
                continue
            if line_str == "Process Group :":
                state = "process"
                continue
            if line_str == "Task/K. Area :":
                state = "task_area"
                continue
            if line_str == "Justification:":
                state = "justification"
                continue

            if line_str.startswith("Domain:"):
                state = "domain"
                domain.append(line_str.split(":", 1)[1].strip())
                continue
            if line_str.startswith("Task:"):
                state = "task_domain"
                task_domain.append(line_str.split(":", 1)[1].strip())
                continue
            if line_str.startswith("Reference:"):
                state = "reference"
                reference.append(line_str.split(":", 1)[1].strip())
                continue

            if state == "choice":
                choices.append(line_str)
                state = "await"
            elif state == "selected":
                selected.append(line_str)
            elif state == "process":
                process_group.append(line_str)
            elif state == "task_area":
                task_area.append(line_str)
            elif state == "justification":
                justification.append(line)
            elif state == "domain":
                domain.append(line_str)
            elif state == "task_domain":
                task_domain.append(line_str)
            elif state == "reference":
                reference.append(line_str)
            elif state == "question":
                question_text.append(line)

        results.append({
            "QuestionNo":    q_no,
            "QID":           q_id,
            "Question":      "\n".join(question_text).strip(),
            "Choices":       "; ".join([c for c in choices if c]),
            "SelectedChoice":  " ".join(selected).strip(),
            "ProcessGroup":  " ".join(process_group).strip(),
            "TaskArea":      " ".join(task_area).strip(),
            "Justification": "\n".join(justification).strip(),
            "Domain":        " ".join(domain).strip(),
            "Task":          " ".join(task_domain).strip(),
            "Reference":     " ".join(reference).strip(),
        })

    return results


def build_markdown(data):
    lines = []
    for entry in data:
        lines.append(f"Question No. :{entry.get('QuestionNo','')}")
        lines.append(f"Q - ID : {entry.get('QID','')}")
        lines.append("")
        lines.append("Question: ")
        lines.append("")
        lines.append(entry.get("Question", ""))
        lines.append("")
        lines.append("Answer: options")

        choices = entry.get("Choices", "").split("; ")
        for idx, choice in enumerate(choices, start=1):
            lines.append(f"Choice {idx} : {choice}")

        lines.append("")
        lines.append(f"Selected Choice: {entry.get('SelectedChoice','')}")
        lines.append("")
        lines.append(f"Process Group : {entry.get('ProcessGroup','')}")
        lines.append("")
        lines.append(f"Task/K. Area : {entry.get('TaskArea','')}")
        lines.append("")

        just_text = entry.get("Justification", "").strip().lstrip('"').rstrip('"')
        ref_text  = entry.get("Reference", "").strip().rstrip(';"').rstrip()

        lines.append("Justification:")
        lines.append(f'"{just_text}')
        lines.append("")
        if entry.get("Domain"):
            lines.append(f"Domain: {entry.get('Domain','')};")
        if entry.get("Task"):
            lines.append(f"Task: {entry.get('Task','')};")
        if entry.get("Reference"):
            lines.append(f'Reference: {ref_text} ;"')
        lines.append("")
        lines.append("")

    return "\n".join(lines)


# ─── DOCX Generator ───────────────────────────────────────────────────────────
# Layout: landscape, narrow margins (0.5"), 2 question boxes side by side per page.
# Landscape content width: 15840 - 720 - 720 = 14400 DXA
# Each question box: (14400 - 360 spacer) / 2 = 7020 DXA
_D_Q_BOX  = 7020
_D_SPACER = 360
_D_LBL    = 1500
_D_VAL    = _D_Q_BOX - _D_LBL   # 5520 DXA
_D_FONT   = 9                    # pt


def _d_fix_tbl_grid(table, col_widths_dxa):
    """Replace auto-generated tblGrid with exact column widths, and set tblW."""
    tbl = table._tbl
    # Fix tblGrid
    for existing in tbl.findall(qn('w:tblGrid')):
        tbl.remove(existing)
    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths_dxa:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    # Insert tblGrid right after tblPr
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)
    # Fix tblW to total fixed width
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        for existing in tblPr.findall(qn('w:tblW')):
            tblPr.remove(existing)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'dxa')
        tblW.set(qn('w:w'), str(sum(col_widths_dxa)))
        tblPr.append(tblW)


def _d_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')):
        tcPr.remove(s)
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:color'), 'auto')
    el.set(qn('w:fill'), fill)
    tcPr.append(el)


def _d_clear_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for ex in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(ex)
    tb = OxmlElement('w:tcBorders')
    for side in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        tb.append(el)
    tcPr.append(tb)


def _d_text(cell, text, bold=False, size=_D_FONT):
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0.5)
    para.paragraph_format.space_after  = Pt(0.5)
    if '\n' in text:
        first = True
        for line in text.split('\n'):
            if not first:
                para.add_run().add_break()
            r = para.add_run(line)
            r.bold = bold
            r.font.color.rgb = RGBColor(0, 0, 0)
            r.font.size = Pt(size)
            first = False
    else:
        r = para.add_run(text)
        r.bold = bold
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt(size)


def _d_text_lb(cell, top, bottom, size=_D_FONT):
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0.5)
    para.paragraph_format.space_after  = Pt(0.5)
    for i, txt in enumerate((top, bottom)):
        if i > 0:
            para.add_run().add_break()
        r = para.add_run(txt)
        r.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt(size)


def _d_move_into_cell(inner_table, outer_cell):
    tc = outer_cell._tc
    for p in list(tc.findall(qn('w:p'))):
        tc.remove(p)
    tc.append(inner_table._tbl)   # lxml re-parents automatically
    tc.append(OxmlElement('w:p')) # OOXML requires trailing paragraph in cell


def _d_build_question_table(doc, entry):
    choices = entry.get("Choices", "").split("; ")
    row_defs = [
        ("Question No",   entry.get("QuestionNo", ""),     "DDD9C3", False, None),
        ("QID",           entry.get("QID", ""),            None,     False, None),
        ("Question",      entry.get("Question", ""),       "F2F2F2", True,  None),
    ]
    for i in range(6):
        row_defs.append((f"Choice {i+1}", choices[i] if i < len(choices) else "", None, False, None))
    row_defs += [
        ("Correct Answer", entry.get("SelectedChoice", ""), "F2F2F2", True,  None),
        (None,             entry.get("ProcessGroup", ""),   None,     False, ("Process", "Group")),
        ("Task Area",      entry.get("TaskArea", ""),       None,     False, None),
        ("Justification",  entry.get("Justification", ""),  "DBE5F1", False, None),
        ("Domain",         entry.get("Domain", ""),         None,     False, None),
        ("Task",           entry.get("Task", ""),           None,     False, None),
        ("Reference",      entry.get("Reference", ""),      None,     False, None),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    for label, value, bg, vbold, lb in row_defs:
        row = tbl.add_row()
        lc, vc = row.cells[0], row.cells[1]
        lc.width = Twips(_D_LBL)
        vc.width = Twips(_D_VAL)
        if bg:
            _d_shading(lc, bg)
            _d_shading(vc, bg)
        if lb:
            _d_text_lb(lc, lb[0], lb[1])
        else:
            _d_text(lc, label or "", bold=True)
        _d_text(vc, value or "", bold=vbold)
    # Fix tblGrid and tblW so Word renders the exact column widths
    _d_fix_tbl_grid(tbl, [_D_LBL, _D_VAL])
    return tbl


def generate_docx_bytes(data, module_name="Output"):
    from datetime import datetime
    doc = Document()

    # Landscape, narrow margins (0.5" = 720 DXA each side)
    sec = doc.sections[0]
    sec.orientation   = WD_ORIENT.LANDSCAPE
    sec.page_width    = Twips(15840)   # 11"
    sec.page_height   = Twips(12240)   # 8.5"
    sec.top_margin    = Twips(720)
    sec.bottom_margin = Twips(720)
    sec.left_margin   = Twips(720)
    sec.right_margin  = Twips(720)

    # Header: module name | conversion date
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(f"{module_name}  |  {datetime.now().strftime('%Y-%m-%d')}")
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(80, 80, 80)

    # Footer: centred page number
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run()
    fr.font.size = Pt(9)
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr  = OxmlElement('w:instrText'); instr.text = 'PAGE'
    end    = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    fr._r.extend([begin, instr, end])

    # Two questions per page, side by side
    for page_i, idx in enumerate(range(0, len(data), 2)):
        if page_i > 0:
            doc.add_page_break()
        left_entry  = data[idx]
        right_entry = data[idx + 1] if idx + 1 < len(data) else None

        # Outer 3-col layout frame: [left_box | spacer | right_box]
        outer = doc.add_table(rows=1, cols=3)
        r = outer.rows[0]
        r.cells[0].width = Twips(_D_Q_BOX)
        r.cells[1].width = Twips(_D_SPACER)
        r.cells[2].width = Twips(_D_Q_BOX)
        for cell in r.cells:
            _d_clear_borders(cell)
        # Fix outer tblGrid so Word respects the 7020/360/7020 column split
        _d_fix_tbl_grid(outer, [_D_Q_BOX, _D_SPACER, _D_Q_BOX])

        left_tbl = _d_build_question_table(doc, left_entry)
        _d_move_into_cell(left_tbl, r.cells[0])

        if right_entry:
            right_tbl = _d_build_question_table(doc, right_entry)
            _d_move_into_cell(right_tbl, r.cells[2])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── PDF Generator ────────────────────────────────────────────────────────────

_COL_TAN  = HexColor('#DDD9C3')
_COL_GRAY = HexColor('#F2F2F2')
_COL_BLUE = HexColor('#DBE5F1')

_BASE = ParagraphStyle('base', fontName='Helvetica',      fontSize=11, leading=15, textColor=colors.black)
_BOLD = ParagraphStyle('bold', fontName='Helvetica-Bold', fontSize=11, leading=15, textColor=colors.black)


def _para(text, bold=False):
    style = _BOLD if bold else _BASE
    safe = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    safe = safe.replace('\n', '<br/>')
    return Paragraph(safe, style)


def generate_pdf_bytes(data):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        topMargin=inch,
        bottomMargin=inch,
        leftMargin=1.25 * inch,
        rightMargin=1.25 * inch,
    )

    avail_w = 6 * inch          # 8.5 - 1.25 - 1.25
    label_w = (2340 / 10100) * avail_w
    value_w = (7760 / 10100) * avail_w

    story = []

    for i, entry in enumerate(data):
        if i > 0:
            story.append(PageBreak())

        choices = entry.get("Choices", "").split("; ")

        rows = [
            [_para("Question No", bold=True), _para(entry.get("QuestionNo", ""))],
            [_para("QID",         bold=True), _para(entry.get("QID", ""))],
            [_para("Question",    bold=True), _para(entry.get("Question", ""), bold=True)],
        ]
        for j in range(6):
            rows.append([_para(f"Choice {j+1}", bold=True), _para(choices[j] if j < len(choices) else "")])
        rows += [
            [_para("Correct Answer", bold=True), _para(entry.get("SelectedChoice", ""), bold=True)],
            [_para("Process\nGroup", bold=True),  _para(entry.get("ProcessGroup", ""))],
            [_para("Task Area",      bold=True),  _para(entry.get("TaskArea", ""))],
            [_para("Justification",  bold=True),  _para(entry.get("Justification", ""))],
            [_para("Domain",         bold=True),  _para(entry.get("Domain", ""))],
            [_para("Task",           bold=True),  _para(entry.get("Task", ""))],
            [_para("Reference",      bold=True),  _para(entry.get("Reference", ""))],
        ]

        t = Table(rows, colWidths=[label_w, value_w])
        # Row indices: QNo=0 QID=1 Q=2 C1-C6=3-8 Ans=9 Proc=10 TA=11 Just=12 Dom=13 Task=14 Ref=15
        t.setStyle(TableStyle([
            ('GRID',          (0, 0),  (-1, -1), 0.5, colors.black),
            ('VALIGN',        (0, 0),  (-1, -1), 'TOP'),
            ('TOPPADDING',    (0, 0),  (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0),  (-1, -1), 5),
            ('LEFTPADDING',   (0, 0),  (-1, -1), 6),
            ('RIGHTPADDING',  (0, 0),  (-1, -1), 6),
            # Row 0 – Question No: tan
            ('BACKGROUND', (0, 0),  (1, 0),  _COL_TAN),
            # Row 2 – Question: light gray
            ('BACKGROUND', (0, 2),  (1, 2),  _COL_GRAY),
            # Row 9 – Correct Answer: light gray
            ('BACKGROUND', (0, 9),  (1, 9),  _COL_GRAY),
            # Row 12 – Justification: light blue
            ('BACKGROUND', (0, 12), (1, 12), _COL_BLUE),
        ]))
        story.append(t)

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ─── Routes ───────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/convert", methods=["POST"])
def convert():
    uploaded = request.files.get("md_file")
    module_no = request.form.get("module_no", "Module").strip()
    export_format = request.form.get("export_format", "md").strip()

    if not uploaded or not uploaded.filename.endswith(".md"):
        return jsonify({"error": "Please upload a valid .md file."}), 400

    content = uploaded.read().decode("utf-8")
    data = parse_markdown(content)

    if not data:
        return jsonify({"error": "No questions found in the uploaded file."}), 400

    if export_format == "csv":
        si = io.StringIO()
        fieldnames = [
            "QuestionNo", "QID", "Question",
            "Choice 1", "Choice 2", "Choice 3", "Choice 4", "Choice 5", "Choice 6",
            "SelectedChoice", "ProcessGroup", "TaskArea", "Justification", "Domain", "Task", "Reference"
        ]
        writer = csv.DictWriter(si, fieldnames=fieldnames)
        writer.writeheader()
        csv_data = []
        for row in data:
            row_copy = dict(row)
            choices_str = row_copy.pop("Choices", "")
            choices_list = choices_str.split("; ") if choices_str else []
            for i in range(1, 7):
                row_copy[f"Choice {i}"] = choices_list[i-1] if i <= len(choices_list) else ""
            csv_data.append(row_copy)
        writer.writerows(csv_data)
        output_bytes = si.getvalue().encode('utf-8-sig')
        mimetype = "text/csv"
        filename = f"{module_no}_output.csv"

    elif export_format == "docx":
        output_bytes = generate_docx_bytes(data, module_name=module_no)
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{module_no}_output.docx"

    elif export_format == "pdf":
        output_bytes = generate_pdf_bytes(data)
        mimetype = "application/pdf"
        filename = f"{module_no}_output.pdf"

    else:  # md
        output_md = build_markdown(data)
        output_bytes = output_md.encode("utf-8")
        mimetype = "text/markdown"
        filename = f"{module_no}_output.md"

    return send_file(
        io.BytesIO(output_bytes),
        mimetype=mimetype,
        as_attachment=True,
        download_name=filename,
    )


@app.route("/shutdown", methods=["POST"])
def shutdown():
    os.kill(os.getpid(), signal.SIGINT)
    return jsonify({"success": True, "message": "Server shutting down..."})


if __name__ == "__main__":
    app.run(host="0.0.0.0", debug=True, port=5000)
